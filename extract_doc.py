# 1. Download the PDF file from the URL.
# 2. Open the PDF file in Python.
# 3. Extract the text from the PDF file.
# 4. Clean up the text.
# 5. Return the cleaned text.
import requests
import os
import pdftotext
import re
from docx import Document
from summarize import gen_summary

def dl_pdf_doc_from_url(url, filename):
    # returns data from url for pdfs
    r = requests.get(url, allow_redirects=True)
    print(r.headers)
    if r.headers["Content-Type"] not in ["application/pdf", "application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        return None
    with open(filename, 'wb') as f:
        f.write(r.content)
    
    return r.content


def post_image_to_discord(url: str, file: str, filename: str = 'file'):
  url = url
  result = requests.post(
      url, files={filename: file}
  )

  try:
      result.raise_for_status()
  except requests.exceptions.HTTPError as err:
      print(err)
  else:
      print("Image Delivered {}.".format(result.status_code))

def mk_dir(new_dir):
    if not os.path.isdir(new_dir):
        os.mkdir(new_dir)

def handle_logic(dataDict = {
    "url": None,
    "path": None,
    "company_name": None
}):
    content = dl_pdf_doc_from_url(dataDict.get("url"), dataDict.get("path"))
    if content == None:
        return {
            "content": None,
            "summary": None,
        }
    pdf_text = extract_from_file(dataDict.get("path"))
    pdf_text = clean_text(pdf_text, dataDict.get("company_name"))
    summary_text = gen_summary(pdf_text)
    file_parts = dataDict.get("path").split(".")
    txt_path = "".join(file_parts[:-1] + [".txt"])
    # save file as text

    with open(txt_path, "w") as f:
        f.write(pdf_text)
    
    return {
        "content": content,
        "summary": summary_text,
    }

def extract_from_file(path: str):
    # try pdf, if that does not work
    # try docx
    # if there are 3 formats, redo this format
    try:
        with open(path, "rb") as f:
            pdf = pdftotext.PDF(f)

        # Read all the text into one string
        pdf_data = "\n\n".join(pdf)
        return pdf_data
    except Exception as e:
        print("FOUND A WORD DOCX")
        # assume word doc
        with open(path, "rb") as doc_file:
            document = Document(doc_file)
        try:
            paragraphs = [x.text for x in document.paragraphs]
            return "\n".join(paragraphs)
        except Exception as e:
            print(e)
            return "\n".join(document.paragraphs)

def clean_text(summary_text: str, company_name: str):
    # remove all text below
    # About {company_name}
    test_phrase = f"About {company_name}"
    try:
        summary_content = summary_text[:summary_text.index(test_phrase)]
        if len(summary_content) > 10:
            return summary_content
        else:
            raise Exception("Please work")
    except Exception as e:
        # try upper case
        print(summary_text)
        updated_company_name = "ABOUT"
        try:
            return summary_text[:summary_text.index(updated_company_name)]
        except Exception as e:
            try:
                print(summary_text)
                search = re.search(summary_text, company_name(), re.IGNORECASE)
                print(search)
                if search:
                    start_index = search.start()
                    return summary_text[:start_index]
                else:
                    raise Exception("I done goofed")
            except Exception as e:
                print(e)
                print("Code is confused, dont judge me, I just want to make money on stonks.")
                return summary_text

            


# handle_logic({"url": "https://webfiles.thecse.com/Peak_Fintech_Continues_Expansion_of_Business_Hub_with_Addition_of_Two_New_Banks_and_New_Office_in_Guangzhou.pdf?dcWC2NW_GzmSFbHOadxyHNESnqSpHkza", "path": "docs/PKK_Peak_Fintech_Continues_Expansion_of_Business_Hub_with_Addition_of_Two_New_Banks_and_New_Office_in_Guangzhou.pdf?dcWC2NW_GzmSFbHOadxyHNESnqSpHkza.pdf"})
